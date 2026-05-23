from __future__ import annotations

from dataclasses import asdict, dataclass
from pathlib import Path
import copy
import json
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


CITE_RE = re.compile(r"\[(\d+)\]")
REF_LINE_RE = re.compile(r"^\s*\[([0-9]+)\]\s*")
BOOKMARK_PREFIX = "ref_"


@dataclass
class ReferenceLinkResult:
    input_path: str
    output_path: str
    bookmarks: dict[str, str]
    linked_citations: int
    skipped_citations: list[str]

    def to_dict(self) -> dict[str, object]:
        return asdict(self)


def _paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def _is_reference_heading(text: str) -> bool:
    return text.replace(" ", "") in {"参考文献", "References", "REFERENCE"}


def _next_bookmark_id(document: Document) -> int:
    values: list[int] = []
    for node in document.element.iter():
        value = node.get(qn("w:id"))
        if value and value.isdigit():
            values.append(int(value))
    return max(values, default=0) + 1


def _remove_existing_reference_bookmarks(document: Document) -> None:
    ids: set[str] = set()
    for node in list(document.element.iter()):
        if node.tag == qn("w:bookmarkStart"):
            name = node.get(qn("w:name")) or ""
            if name.startswith(BOOKMARK_PREFIX):
                bookmark_id = node.get(qn("w:id"))
                if bookmark_id:
                    ids.add(bookmark_id)
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)
    for node in list(document.element.iter()):
        if node.tag == qn("w:bookmarkEnd") and node.get(qn("w:id")) in ids:
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)


def _insert_reference_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    p_element = paragraph._p
    insert_index = 0
    if len(p_element) and p_element[0].tag == qn("w:pPr"):
        insert_index = 1
    p_element.insert(insert_index, start)
    p_element.append(end)


def _collect_reference_paragraphs(document: Document) -> dict[str, object]:
    references: dict[str, object] = {}
    in_references = False
    for paragraph in document.paragraphs:
        text = _paragraph_text(paragraph)
        if _is_reference_heading(text):
            in_references = True
            continue
        if not in_references:
            continue
        match = REF_LINE_RE.match(text)
        if match:
            references[match.group(1)] = paragraph
    return references


def _text_run_like(source_run, text: str):
    run = OxmlElement("w:r")
    rpr = source_run.find(qn("w:rPr"))
    if rpr is not None:
        run.append(copy.deepcopy(rpr))
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    node.text = text
    run.append(node)
    return run


def _hyperlink_run_like(source_run, text: str, anchor: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = source_run.find(qn("w:rPr"))
    if rpr is not None:
        rpr = copy.deepcopy(rpr)
    else:
        rpr = OxmlElement("w:rPr")

    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.insert(0, rstyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)

    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    return hyperlink


def _link_paragraph_citations(paragraph, available_refs: set[str]) -> tuple[int, list[str]]:
    p_element = paragraph._p
    linked = 0
    skipped: list[str] = []
    for child in list(p_element):
        if child.tag != qn("w:r"):
            continue
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t")))
        if not text or not CITE_RE.search(text):
            continue
        pieces = []
        pos = 0
        changed = False
        for match in CITE_RE.finditer(text):
            if match.start() > pos:
                pieces.append(("text", text[pos : match.start()], ""))
            number = match.group(1)
            if number in available_refs:
                pieces.append(("link", match.group(0), BOOKMARK_PREFIX + number))
                linked += 1
                changed = True
            else:
                pieces.append(("text", match.group(0), ""))
                skipped.append(match.group(0))
            pos = match.end()
        if pos < len(text):
            pieces.append(("text", text[pos:], ""))
        if not changed:
            continue
        insert_at = p_element.index(child)
        p_element.remove(child)
        for kind, value, anchor in pieces:
            if not value:
                continue
            new_node = _hyperlink_run_like(child, value, anchor) if kind == "link" else _text_run_like(child, value)
            p_element.insert(insert_at, new_node)
            insert_at += 1
    return linked, skipped


def link_reference_citations(input_path: str | Path, output_path: str | Path) -> ReferenceLinkResult:
    source = Path(input_path)
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(source))
    _remove_existing_reference_bookmarks(document)
    references = _collect_reference_paragraphs(document)

    next_id = _next_bookmark_id(document)
    bookmarks: dict[str, str] = {}
    for number, paragraph in sorted(references.items(), key=lambda item: int(item[0])):
        name = BOOKMARK_PREFIX + number
        _insert_reference_bookmark(paragraph, name, next_id)
        next_id += 1
        bookmarks[number] = name

    linked = 0
    skipped: list[str] = []
    in_references = False
    for paragraph in document.paragraphs:
        text = _paragraph_text(paragraph)
        if _is_reference_heading(text):
            in_references = True
        if in_references:
            continue
        count, missing = _link_paragraph_citations(paragraph, set(references))
        linked += count
        skipped.extend(missing)

    document.save(str(target))
    return ReferenceLinkResult(str(source), str(target), bookmarks, linked, skipped)


def write_reference_link_report(result: ReferenceLinkResult, output_path: str | Path) -> Path:
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(result.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    return target
