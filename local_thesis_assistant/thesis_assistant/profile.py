from __future__ import annotations

from dataclasses import asdict, dataclass, field
import json
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from .doc_reader import ensure_docx
from .rules import ThesisRules, default_rules, save_rules


@dataclass
class TemplateStyle:
    name: str
    type: str
    font_east_asia: str = ""
    font_ascii: str = ""
    font_size_pt: float | None = None
    bold: bool | None = None


@dataclass
class TemplateProfile:
    name: str
    source_path: str
    docx_path: str
    standard_base: str = "GB/T 7713.1-2025 + GB/T 7714-2015"
    reference_standard: str = "GB/T 7714-2015"
    rules: dict[str, Any] = field(default_factory=dict)
    styles: list[TemplateStyle] = field(default_factory=list)
    page: dict[str, float] = field(default_factory=dict)
    counts: dict[str, int] = field(default_factory=dict)
    notes: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        data["styles"] = [asdict(item) for item in self.styles]
        return data

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "TemplateProfile":
        styles = [TemplateStyle(**item) for item in data.get("styles", [])]
        return cls(
            name=data.get("name", "template-profile"),
            source_path=data.get("source_path", ""),
            docx_path=data.get("docx_path", ""),
            standard_base=data.get("standard_base", "GB/T 7713.1-2025 + GB/T 7714-2015"),
            reference_standard=data.get("reference_standard", "GB/T 7714-2015"),
            rules=data.get("rules", {}),
            styles=styles,
            page=data.get("page", {}),
            counts=data.get("counts", {}),
            notes=data.get("notes", []),
        )


def _style_font(style) -> tuple[str, str, float | None, bool | None]:
    font = style.font
    east = ""
    ascii_font = font.name or ""
    size = round(font.size.pt, 1) if font.size is not None else None
    bold = font.bold
    rpr = style._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        east = rpr.rFonts.get(qn("w:eastAsia")) or ""
    return east, ascii_font, size, bold


def _run_east_font(run) -> str:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return ""
    return rpr.rFonts.get(qn("w:eastAsia")) or ""


def _dominant_body_size(document: Document) -> float | None:
    values: Counter[float] = Counter()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text.startswith(("图", "表", "[")):
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            continue
        for run in paragraph.runs:
            if run.font.size is not None:
                values[round(run.font.size.pt, 1)] += 1
    if values:
        return values.most_common(1)[0][0]
    return None


def _dominant_body_font(document: Document) -> str:
    values: Counter[str] = Counter()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text.startswith(("图", "表", "[")):
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            continue
        for run in paragraph.runs:
            east = _run_east_font(run)
            if east:
                values[east] += 1
    if values:
        return values.most_common(1)[0][0]
    return ""


def inspect_template(template_path: str | Path, work_dir: str | Path | None = None) -> TemplateProfile:
    source = Path(template_path)
    try:
        docx_path = ensure_docx(source, work_dir or source.parent)
    except Exception as exc:
        rules = default_rules()
        return TemplateProfile(
            name=source.stem,
            source_path=str(source),
            docx_path="",
            rules=rules.to_dict(),
            styles=[],
            page={
                "top_margin_cm": rules.margins.top_cm,
                "bottom_margin_cm": rules.margins.bottom_cm,
                "left_margin_cm": rules.margins.left_cm,
                "right_margin_cm": rules.margins.right_cm,
            },
            counts={"paragraphs": 0, "tables": 0, "images": 0, "sections": 0},
            notes=[
                "模板转换失败，已生成通用本科论文规则画像作为降级结果。",
                f"转换诊断：{exc}",
                "建议用 Microsoft Word 手动将学校 .doc 模板另存为 .docx 后重新运行 template-inspect。",
            ],
        )
    document = Document(str(docx_path))
    rules = default_rules()

    if document.sections:
        section = document.sections[0]
        rules.margins.top_cm = round(section.top_margin.cm, 2)
        rules.margins.bottom_cm = round(section.bottom_margin.cm, 2)
        rules.margins.left_cm = round(section.left_margin.cm, 2)
        rules.margins.right_cm = round(section.right_margin.cm, 2)
        page = {
            "width_pt": round(section.page_width.pt, 2),
            "height_pt": round(section.page_height.pt, 2),
            "top_margin_cm": rules.margins.top_cm,
            "bottom_margin_cm": rules.margins.bottom_cm,
            "left_margin_cm": rules.margins.left_cm,
            "right_margin_cm": rules.margins.right_cm,
        }
    else:
        page = {}

    body_size = _dominant_body_size(document)
    body_font = _dominant_body_font(document)
    if body_size:
        rules.body.font_size_pt = body_size
    if body_font:
        rules.body.font_east_asia = body_font

    styles: list[TemplateStyle] = []
    for style in document.styles:
        east, ascii_font, size, bold = _style_font(style)
        styles.append(
            TemplateStyle(
                name=style.name,
                type=str(style.type),
                font_east_asia=east,
                font_ascii=ascii_font,
                font_size_pt=size,
                bold=bold,
            )
        )

    picture_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    profile = TemplateProfile(
        name=source.stem,
        source_path=str(source),
        docx_path=str(docx_path),
        rules=rules.to_dict(),
        styles=styles,
        page=page,
        counts={
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "images": picture_count,
            "sections": len(document.sections),
        },
        notes=[
            "学校模板优先；无法从模板明确抽取的项目使用通用本科论文基线。",
            "参考文献默认使用 GB/T 7714-2015，并预留 GB/T 7714-2025 切换。",
        ],
    )
    return profile


def load_profile(path: str | Path | None) -> TemplateProfile:
    if path is None:
        rules = default_rules()
        return TemplateProfile(
            name="default-undergraduate-profile",
            source_path="",
            docx_path="",
            rules=rules.to_dict(),
            notes=["未提供模板画像，使用通用本科论文基线。"],
        )
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    return TemplateProfile.from_dict(data)


def profile_rules(profile: TemplateProfile) -> ThesisRules:
    if profile.rules:
        return ThesisRules.from_dict(profile.rules)
    return default_rules()


def write_profile(profile: TemplateProfile, output_path: str | Path) -> Path:
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(profile.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def write_profile_rules(profile: TemplateProfile, output_path: str | Path) -> Path:
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    save_rules(profile_rules(profile), target)
    return target
