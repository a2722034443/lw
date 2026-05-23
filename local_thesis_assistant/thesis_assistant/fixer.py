from __future__ import annotations

from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .formatter import normalize_docx
from .profile import TemplateProfile, load_profile, profile_rules
from .rules import ThesisRules


def _set_cell_font(table, east_asia: str, ascii_font: str, size_pt: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = ascii_font
                    run.font.size = Pt(size_pt)
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.get_or_add_rFonts()
                    rfonts.set(qn("w:eastAsia"), east_asia)


def _border(name: str, val: str, size: str = "8") -> OxmlElement:
    item = OxmlElement(f"w:{name}")
    item.set(qn("w:val"), val)
    if val not in {"nil", "none"}:
        item.set(qn("w:sz"), size)
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "000000")
    return item


def _set_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    borders.append(_border("top", "single", "12"))
    borders.append(_border("left", "nil"))
    borders.append(_border("bottom", "single", "12"))
    borders.append(_border("right", "nil"))
    borders.append(_border("insideH", "single", "6"))
    borders.append(_border("insideV", "nil"))
    tbl_pr.append(borders)


def _rewrite_zip_member(path: Path, member: str, content: str) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp:
        temp_path = Path(temp.name)
    with zipfile.ZipFile(path, "r") as src, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            if info.filename == member:
                dst.writestr(info.filename, content)
            else:
                dst.writestr(info, src.read(info.filename))
    shutil.move(str(temp_path), path)


def _set_update_fields_on_open(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as archive:
        if "word/settings.xml" not in archive.namelist():
            return
        xml = archive.read("word/settings.xml").decode("utf-8", errors="ignore")
    if "w:updateFields" in xml:
        xml = xml.replace('w:updateFields w:val="false"', 'w:updateFields w:val="true"')
        xml = xml.replace('w:updateFields w:val="0"', 'w:updateFields w:val="true"')
    else:
        xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
    _rewrite_zip_member(path, "word/settings.xml", xml)


def _normalize_tables(path: Path, rules: ThesisRules) -> None:
    document = Document(str(path))
    for table in document.tables:
        _set_three_line_table(table)
        _set_cell_font(table, rules.body.font_east_asia, rules.body.font_ascii, rules.body.font_size_pt)
    document.save(str(path))


def conservative_fix_docx(
    input_path: str | Path,
    output_path: str | Path,
    profile_path: str | Path | None = None,
    profile: TemplateProfile | None = None,
) -> Path:
    active_profile = profile or load_profile(profile_path)
    rules = profile_rules(active_profile)
    target = normalize_docx(input_path, output_path, rules)
    _normalize_tables(target, rules)
    _set_update_fields_on_open(target)
    return target
