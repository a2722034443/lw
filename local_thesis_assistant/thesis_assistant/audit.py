from __future__ import annotations

from dataclasses import asdict, dataclass
import json
import re
import zipfile
from collections import defaultdict
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .doc_reader import ensure_docx
from .paragraph_ooxml import paragraph_ooxml_summary
from .profile import TemplateProfile, load_profile, profile_rules
from .rules import ThesisRules


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "a": A_NS, "wp": WP_NS, "m": M_NS}

ALIGNMENT_NAME = {
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
}
HEADING_RE = re.compile(r"^第([一二三四五六七八九十\d]+)章")
SECTION_RE = re.compile(r"^(\d+)\.(\d+)\s+")
SUBSECTION_RE = re.compile(r"^(\d+)\.(\d+)\.(\d+)\s+")
FIGURE_RE = re.compile(r"^\s*图\s*([0-9]+)[\.-]([0-9]+)")
TABLE_RE = re.compile(r"^\s*表\s*([0-9]+)[\.-]([0-9]+)")
FORMULA_CAPTION_RE = re.compile(r"[（(]\s*([0-9]+)[\.-]([0-9]+)\s*[）)]\s*$")
CITE_RE = re.compile(r"\[([0-9]+)\]")
REF_LINE_RE = re.compile(r"^\s*\[([0-9]+)\]\s*(.+)")


@dataclass
class Finding:
    code: str
    severity: str
    location: str
    message: str
    expected: str
    actual: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass
class AuditResult:
    file: str
    profile: str
    findings: list[Finding]
    metrics: dict[str, object]

    def to_dict(self) -> dict[str, object]:
        return {
            "file": self.file,
            "profile": self.profile,
            "findings": [item.to_dict() for item in self.findings],
            "metrics": self.metrics,
        }

    @property
    def has_errors(self) -> bool:
        return any(item.severity == "error" for item in self.findings)


def _cm(value) -> float:
    return round(value.cm, 2)


def _pt(value) -> float | None:
    if value is None:
        return None
    return round(value.pt, 1)


def _run_east_font(run) -> str:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return ""
    return rpr.rFonts.get(qn("w:eastAsia")) or ""


def _style_east_font(paragraph) -> str:
    if not paragraph.style:
        return ""
    rpr = paragraph.style._element.rPr
    if rpr is None or rpr.rFonts is None:
        return ""
    return rpr.rFonts.get(qn("w:eastAsia")) or ""


def _effective_font_size(paragraph, run) -> float | None:
    if run.font.size is not None:
        return round(run.font.size.pt, 1)
    if paragraph.style and paragraph.style.font.size is not None:
        return round(paragraph.style.font.size.pt, 1)
    return None


def _effective_east_font(paragraph, run) -> str:
    return _run_east_font(run) or _style_east_font(paragraph) or run.font.name or ""


def _paragraph_level(paragraph) -> int | None:
    text = paragraph.text.strip()
    if paragraph.style and paragraph.style.name.lower().startswith("heading"):
        match = re.search(r"(\d+)", paragraph.style.name)
        if match:
            return int(match.group(1))
    if HEADING_RE.match(text):
        return 1
    if SUBSECTION_RE.match(text):
        return 3
    if SECTION_RE.match(text):
        return 2
    return None


def _paragraph_has_image(paragraph) -> bool:
    xml = paragraph._element.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8", errors="ignore")


def _field_codes(xml: str) -> list[str]:
    root = ET.fromstring(xml)
    codes: list[str] = []
    for node in root.findall(".//w:instrText", NS):
        if node.text:
            codes.append(" ".join(node.text.split()))
    return codes


def _bookmark_names(xml: str) -> set[str]:
    root = ET.fromstring(xml)
    names: set[str] = set()
    for node in root.findall(".//w:bookmarkStart", NS):
        name = node.get(f"{{{W_NS}}}name")
        if name:
            names.add(name)
    return names


def _hyperlink_anchors(xml: str) -> list[str]:
    root = ET.fromstring(xml)
    anchors: list[str] = []
    for node in root.findall(".//w:hyperlink", NS):
        anchor = node.get(f"{{{W_NS}}}anchor")
        if anchor:
            anchors.append(anchor)
    return anchors


def _omml_count(xml: str) -> int:
    root = ET.fromstring(xml)
    return len(root.findall(".//m:oMath", NS)) + len(root.findall(".//m:oMathPara", NS))


def _media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        return len([name for name in archive.namelist() if name.startswith("word/media/")])


def _caption_numbers(texts: Iterable[str], regex: re.Pattern[str]) -> list[tuple[int, int]]:
    numbers: list[tuple[int, int]] = []
    for text in texts:
        match = regex.search(text)
        if match:
            numbers.append((int(match.group(1)), int(match.group(2))))
    return numbers


def _duplicate_or_gap(numbers: list[tuple[int, int]]) -> tuple[list[str], list[str]]:
    by_chapter: dict[int, list[int]] = defaultdict(list)
    for chapter, item in numbers:
        by_chapter[chapter].append(item)
    duplicates: list[str] = []
    gaps: list[str] = []
    for chapter, items in by_chapter.items():
        seen: set[int] = set()
        for item in items:
            if item in seen:
                duplicates.append(f"{chapter}.{item}")
            seen.add(item)
        if seen:
            expected = list(range(min(seen), max(seen) + 1))
            gaps.extend(f"{chapter}.{item}" for item in expected if item not in seen)
    return sorted(duplicates), sorted(gaps)


def _reference_numbers(texts: list[str]) -> list[int]:
    numbers: list[int] = []
    in_refs = False
    for text in texts:
        if text.strip() == "参考文献":
            in_refs = True
            continue
        if not in_refs:
            continue
        match = REF_LINE_RE.match(text)
        if match:
            numbers.append(int(match.group(1)))
    return numbers


def _first_seen_citations(texts: list[str]) -> list[int]:
    seen: list[int] = []
    in_refs = False
    for text in texts:
        if text.strip() == "参考文献":
            in_refs = True
        if in_refs:
            continue
        for raw in CITE_RE.findall(text):
            number = int(raw)
            if number not in seen:
                seen.append(number)
    return seen


def _iter_blocks(document: Document):
    paragraph_map = {paragraph._p: (index, paragraph) for index, paragraph in enumerate(document.paragraphs)}
    table_map = {table._tbl: (index, table) for index, table in enumerate(document.tables)}
    for child in document.element.body.iterchildren():
        if child in paragraph_map:
            index, paragraph = paragraph_map[child]
            yield "paragraph", index, paragraph
        elif child in table_map:
            index, table = table_map[child]
            yield "table", index, table


def _table_border_val(table, name: str) -> str:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return ""
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return ""
    border = borders.find(qn(f"w:{name}"))
    if border is None:
        return ""
    return border.get(qn("w:val")) or ""


def _is_three_line_table(table) -> bool:
    top = _table_border_val(table, "top")
    bottom = _table_border_val(table, "bottom")
    inside_h = _table_border_val(table, "insideH")
    left = _table_border_val(table, "left")
    right = _table_border_val(table, "right")
    inside_v = _table_border_val(table, "insideV")
    has_horizontal = all(value and value != "nil" for value in [top, bottom, inside_h])
    no_vertical = all(value in {"", "nil", "none"} for value in [left, right, inside_v])
    return has_horizontal and no_vertical


def _check_margins(document: Document, rules: ThesisRules, findings: list[Finding]) -> None:
    for index, section in enumerate(document.sections):
        checks = [
            ("top", _cm(section.top_margin), rules.margins.top_cm, "上边距"),
            ("bottom", _cm(section.bottom_margin), rules.margins.bottom_cm, "下边距"),
            ("left", _cm(section.left_margin), rules.margins.left_cm, "左边距"),
            ("right", _cm(section.right_margin), rules.margins.right_cm, "右边距"),
        ]
        for code, actual, expected, label in checks:
            if abs(actual - expected) > 0.08:
                findings.append(Finding(f"page_margin_{code}", "error", f"section[{index}]", f"{label}不符合规则", f"{expected} cm", f"{actual} cm"))


def _check_structure(texts: list[str], findings: list[Finding]) -> None:
    joined = "\n".join(texts)
    compact_joined = re.sub(r"\s+", "", joined)
    required = [("摘要", "摘要"), ("Abstract", "英文摘要"), ("参考文献", "参考文献")]
    for token, label in required:
        haystack = compact_joined if token != "Abstract" else joined
        if token not in haystack:
            findings.append(Finding("structure_missing", "error", "document", f"缺少{label}", label, "未找到"))
    if "致谢" not in joined:
        findings.append(Finding("structure_acknowledgement", "warning", "document", "未检测到致谢部分", "致谢", "未找到"))

    chapters: list[int] = []
    for text in texts:
        match = HEADING_RE.match(text.strip())
        if match:
            raw = match.group(1)
            if raw.isdigit():
                chapters.append(int(raw))
    if chapters:
        expected = list(range(min(chapters), max(chapters) + 1))
        if chapters != expected:
            findings.append(Finding("chapter_sequence", "error", "headings", "章节编号不连续或顺序异常", ",".join(map(str, expected)), ",".join(map(str, chapters))))


def _check_paragraph_format(document: Document, rules: ThesisRules, findings: list[Finding]) -> None:
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        level = _paragraph_level(paragraph)
        is_caption = bool(FIGURE_RE.match(text) or TABLE_RE.match(text))
        if level:
            rule = next((item for item in rules.headings if item.level == min(level, 3)), None)
            if not rule:
                continue
            expected_size = rule.font_size_pt
            expected_font = rule.font_east_asia
            expected_alignment = rule.alignment
            expected_message = "标题"
        elif is_caption:
            expected_size = rules.caption.font_size_pt
            expected_font = rules.caption.font_east_asia
            expected_alignment = rules.caption.alignment
            expected_message = "题注"
        else:
            if text == "参考文献" or REF_LINE_RE.match(text):
                continue
            expected_size = rules.body.font_size_pt
            expected_font = rules.body.font_east_asia
            expected_alignment = rules.body.alignment
            expected_message = "正文"
            indent = _pt(paragraph.paragraph_format.first_line_indent)
            if indent is not None and abs(indent - 24.0) > 3.0:
                findings.append(Finding("body_first_line_indent", "warning", f"paragraph[{index}]", "正文首行缩进可能不符合规则", "约 2 字符", f"{indent} pt"))
            spacing = paragraph.paragraph_format.line_spacing
            if isinstance(spacing, float) and abs(spacing - rules.body.line_spacing) > 0.1:
                findings.append(Finding("body_line_spacing", "warning", f"paragraph[{index}]", "正文行距不符合规则", str(rules.body.line_spacing), str(spacing)))
            summary = paragraph_ooxml_summary(paragraph)
            expected_first = str(int(round(rules.body.first_line_indent_chars * 100)))
            actual_first = summary.get("first_line_chars")
            if rules.body.special_indent == "first_line" and actual_first not in {None, expected_first}:
                findings.append(Finding("body_first_line_indent_chars", "warning", f"paragraph[{index}]", "正文首行字符缩进不符合规则", expected_first, str(actual_first)))
            expected_before = str(int(round(rules.body.space_before_lines * 100)))
            expected_after = str(int(round(rules.body.space_after_lines * 100)))
            if summary.get("before_lines") not in {None, expected_before}:
                findings.append(Finding("body_space_before_lines", "warning", f"paragraph[{index}]", "正文段前行数不符合规则", expected_before, str(summary.get("before_lines"))))
            if summary.get("after_lines") not in {None, expected_after}:
                findings.append(Finding("body_space_after_lines", "warning", f"paragraph[{index}]", "正文段后行数不符合规则", expected_after, str(summary.get("after_lines"))))
            expected_line = str(int(round(rules.body.line_spacing * 240))) if rules.body.line_spacing_rule == "auto" else str(int(round(rules.body.line_spacing * 20)))
            if summary.get("line") not in {None, expected_line}:
                findings.append(Finding("body_line_spacing_ooxml", "warning", f"paragraph[{index}]", "正文 OOXML 行距不符合规则", expected_line, str(summary.get("line"))))
            for field, label in [
                ("widow_control", "孤行控制"),
                ("keep_with_next", "与下段同页"),
                ("keep_together", "段中不分页"),
                ("page_break_before", "段前分页"),
                ("kinsoku", "按中文习惯控制首尾字符"),
                ("word_wrap", "允许西文在单词中间换行"),
                ("overflow_punctuation", "允许标点溢出边界"),
                ("auto_space_east_asian_latin", "自动调整中文与西文间距"),
                ("auto_space_east_asian_digit", "自动调整中文与数字间距"),
            ]:
                actual = summary.get(field)
                expected = bool(getattr(rules.body, field))
                if actual is not None and actual != expected:
                    findings.append(Finding(f"body_{field}", "warning", f"paragraph[{index}]", f"正文{label}设置不符合规则", str(expected), str(actual)))

        alignment = ALIGNMENT_NAME.get(paragraph.alignment, "inherit")
        if alignment != "inherit" and alignment != expected_alignment:
            findings.append(Finding(f"{expected_message}_alignment", "warning", f"paragraph[{index}]", f"{expected_message}对齐方式不符合规则", expected_alignment, alignment))

        checked_runs = 0
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            checked_runs += 1
            actual_size = _effective_font_size(paragraph, run)
            actual_font = _effective_east_font(paragraph, run)
            if actual_size is not None and abs(actual_size - expected_size) > 0.2:
                findings.append(Finding(f"{expected_message}_font_size", "warning", f"paragraph[{index}]", f"{expected_message}字号不符合规则", f"{expected_size} pt", f"{actual_size} pt"))
                break
            if actual_font and expected_font not in actual_font:
                findings.append(Finding(f"{expected_message}_font", "warning", f"paragraph[{index}]", f"{expected_message}东亚字体不符合规则", expected_font, actual_font))
                break
        if checked_runs == 0 and text:
            findings.append(Finding("paragraph_no_runs", "warning", f"paragraph[{index}]", "段落没有可检查的运行文本", "存在文本运行", text[:20]))


def _check_captions_and_media(document: Document, docx_path: Path, texts: list[str], findings: list[Finding]) -> dict[str, int]:
    figure_numbers = _caption_numbers(texts, FIGURE_RE)
    table_numbers = _caption_numbers(texts, TABLE_RE)
    formula_numbers = _caption_numbers(texts, FORMULA_CAPTION_RE)
    for label, numbers, code in [("图号", figure_numbers, "figure_number"), ("表号", table_numbers, "table_number"), ("公式号", formula_numbers, "formula_number")]:
        duplicates, gaps = _duplicate_or_gap(numbers)
        if duplicates:
            findings.append(Finding(code, "error", "captions", f"{label}存在重复", "不重复", ",".join(duplicates)))
        if gaps:
            findings.append(Finding(code + "_gap", "warning", "captions", f"{label}存在跳号", "按章连续", ",".join(gaps)))

    blocks = list(_iter_blocks(document))
    for pos, (kind, index, item) in enumerate(blocks):
        if kind != "paragraph":
            continue
        text = item.text.strip()
        if FIGURE_RE.match(text):
            previous_has_image = pos > 0 and blocks[pos - 1][0] == "paragraph" and _paragraph_has_image(blocks[pos - 1][2])
            if not previous_has_image:
                findings.append(Finding("figure_caption_position", "warning", f"paragraph[{index}]", "图题前未检测到图片段落", "图片后紧跟图题", text))
        if TABLE_RE.match(text):
            next_is_table = pos + 1 < len(blocks) and blocks[pos + 1][0] == "table"
            if not next_is_table:
                findings.append(Finding("table_caption_position", "warning", f"paragraph[{index}]", "表题后未检测到表格", "表题后紧跟表格", text))

    media = _media_count(docx_path)
    image_paragraphs = sum(1 for paragraph in document.paragraphs if _paragraph_has_image(paragraph))
    if len(figure_numbers) != media:
        findings.append(Finding("figure_media_count", "warning", "document", "图题数量与嵌入媒体数量不一致", f"{media}", f"{len(figure_numbers)}"))
    return {"figure_captions": len(figure_numbers), "table_captions": len(table_numbers), "embedded_media": media, "image_paragraphs": image_paragraphs}


def _check_tables(document: Document, findings: list[Finding]) -> None:
    for index, table in enumerate(document.tables):
        if len(table.rows) < 2 or len(table.columns) < 1:
            findings.append(Finding("table_shape", "warning", f"table[{index}]", "表格行列数量过少", "至少 2 行", f"{len(table.rows)} 行"))
        if not _is_three_line_table(table):
            findings.append(Finding("table_three_line", "warning", f"table[{index}]", "表格可能不是三线表", "顶线/表头线/底线，无竖线", "未匹配"))


def _check_references(texts: list[str], findings: list[Finding]) -> dict[str, object]:
    citations = _first_seen_citations(texts)
    references = _reference_numbers(texts)
    if citations and citations != sorted(citations):
        findings.append(Finding("citation_order", "error", "body", "正文引用首次出现顺序不是递增顺序", "1..N 递增", ",".join(map(str, citations))))
    if references:
        expected = list(range(references[0], references[0] + len(references)))
        if references != expected:
            findings.append(Finding("reference_sequence", "error", "references", "参考文献编号不连续", ",".join(map(str, expected)), ",".join(map(str, references))))
    if citations and references:
        missing = [item for item in citations if item not in references]
        unused = [item for item in references if item not in citations]
        if missing:
            findings.append(Finding("citation_missing_reference", "error", "references", "正文引用在文末不存在", "文末存在", ",".join(map(str, missing))))
        if unused:
            findings.append(Finding("reference_unused", "warning", "references", "参考文献未被正文引用", "正文引用全部文献", ",".join(map(str, unused))))
    if not citations:
        findings.append(Finding("citation_absent", "warning", "body", "正文未检测到顺序编码引用", "形如 [1] 的引用", "未找到"))
    return {"first_seen_citations": citations, "reference_numbers": references}


def _check_reference_links(xml: str, texts: list[str], findings: list[Finding]) -> dict[str, object]:
    references = _reference_numbers(texts)
    citations = _first_seen_citations(texts)
    bookmarks = _bookmark_names(xml)
    anchors = _hyperlink_anchors(xml)
    missing_bookmarks = [str(number) for number in references if f"ref_{number}" not in bookmarks]
    if references and missing_bookmarks:
        findings.append(
            Finding(
                "reference_bookmark_absent",
                "warning",
                "references",
                "参考文献条目缺少可跳转书签",
                "ref_N 书签",
                ",".join(missing_bookmarks),
            )
        )
    missing_links = [str(number) for number in citations if f"ref_{number}" not in anchors]
    if citations and missing_links:
        findings.append(
            Finding(
                "citation_hyperlink_absent",
                "warning",
                "body",
                "正文引用缺少指向参考文献的内部超链接",
                "w:hyperlink w:anchor=ref_N",
                ",".join(missing_links),
            )
        )
    broken = [anchor for anchor in anchors if anchor.startswith("ref_") and anchor not in bookmarks]
    if broken:
        findings.append(Finding("citation_hyperlink_broken", "error", "body", "正文引用超链接指向不存在的书签", "有效 ref_N 书签", ",".join(sorted(set(broken)))))
    return {
        "reference_bookmarks": len([name for name in bookmarks if name.startswith("ref_")]),
        "reference_hyperlinks": len([anchor for anchor in anchors if anchor.startswith("ref_")]),
    }


def _check_formulas(docx_path: Path, texts: list[str], findings: list[Finding]) -> dict[str, int]:
    xml = _document_xml(docx_path)
    formula_count = _omml_count(xml)
    formula_numbers = len(_caption_numbers(texts, FORMULA_CAPTION_RE))
    if formula_count and formula_numbers == 0:
        findings.append(Finding("formula_number_absent", "warning", "formulas", "检测到 OMML 公式但未检测到公式编号", "公式编号", "未找到"))
    if formula_numbers and formula_count == 0:
        findings.append(Finding("formula_object_absent", "warning", "formulas", "检测到公式编号但未检测到 OMML 公式对象", "OMML 公式", "未找到"))
    return {"omml_formulas": formula_count, "formula_captions": formula_numbers}


def audit_document(input_path: str | Path, profile_path: str | Path | None = None, profile: TemplateProfile | None = None) -> AuditResult:
    active_profile = profile or load_profile(profile_path)
    rules = profile_rules(active_profile)
    docx_path = ensure_docx(input_path)
    document = Document(str(docx_path))
    findings: list[Finding] = []
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    _check_margins(document, rules, findings)
    _check_structure(texts, findings)
    _check_paragraph_format(document, rules, findings)
    caption_metrics = _check_captions_and_media(document, docx_path, texts, findings)
    _check_tables(document, findings)
    reference_metrics = _check_references(texts, findings)
    xml = _document_xml(docx_path)
    reference_link_metrics = _check_reference_links(xml, texts, findings)
    formula_metrics = _check_formulas(docx_path, texts, findings)
    codes = _field_codes(xml)

    metrics: dict[str, object] = {
        "paragraphs": len(document.paragraphs),
        "non_empty_paragraphs": len(texts),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "field_codes": {
            "SEQ": len([item for item in codes if re.search(r"\bSEQ\b", item)]),
            "REF_OR_PAGEREF": len([item for item in codes if re.search(r"\bREF\b|\bPAGEREF\b", item)]),
            "TOC": len([item for item in codes if re.search(r"\bTOC\b", item)]),
        },
        "paragraph_rule": rules.body.__dict__,
        **caption_metrics,
        **reference_metrics,
        **reference_link_metrics,
        **formula_metrics,
    }
    return AuditResult(str(docx_path), active_profile.name, findings, metrics)


def write_audit_json(result: AuditResult, output_path: str | Path) -> Path:
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(result.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def write_audit_markdown(result: AuditResult, output_path: str | Path) -> Path:
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    findings = result.findings
    lines = [
        "# DOCX 论文全量审查报告",
        "",
        f"- 文件：{result.file}",
        f"- 模板画像：{result.profile}",
        f"- 问题总数：{len(findings)}",
        f"- 错误：{sum(1 for item in findings if item.severity == 'error')}",
        f"- 警告：{sum(1 for item in findings if item.severity == 'warning')}",
        "",
        "## 指标",
        "",
        "```json",
        json.dumps(result.metrics, ensure_ascii=False, indent=2),
        "```",
        "",
        "## 问题列表",
        "",
        "| 编码 | 级别 | 位置 | 问题 | 期望 | 实际 |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for item in findings:
        lines.append(f"| {item.code} | {item.severity} | {item.location} | {item.message} | {item.expected} | {item.actual} |")
    target.write_text("\n".join(lines), encoding="utf-8")
    return target
