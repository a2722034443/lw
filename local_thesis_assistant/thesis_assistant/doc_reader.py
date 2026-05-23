from __future__ import annotations

from dataclasses import dataclass
import os
from pathlib import Path
import multiprocessing as mp
import shutil
import subprocess
from typing import Iterable

from docx import Document


@dataclass
class ParagraphInfo:
    index: int
    text: str
    style: str


@dataclass
class TableInfo:
    index: int
    rows: int
    cols: int


@dataclass
class DocumentInfo:
    path: Path
    paragraphs: list[ParagraphInfo]
    tables: list[TableInfo]
    picture_count: int
    section_count: int


def can_convert_doc() -> dict[str, bool]:
    try:
        import win32com.client  # type: ignore  # noqa: F401

        word_com = True
    except Exception:
        word_com = False
    return {
        "soffice": shutil.which("soffice") is not None,
        "winword": shutil.which("winword") is not None,
        "word_com": word_com,
    }


def _word_com_convert_worker(source: str, target: str, queue) -> None:
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            doc = word.Documents.Open(
                FileName=source,
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
                OpenAndRepair=True,
                NoEncodingDialog=True,
            )
            try:
                if hasattr(doc, "SaveAs2"):
                    doc.SaveAs2(FileName=target, FileFormat=16)
                else:
                    doc.SaveAs(FileName=target, FileFormat=16)
            finally:
                doc.Close(False)
        finally:
            word.Quit()
        queue.put({"ok": True, "error": ""})
    except Exception as exc:  # pragma: no cover - depends on local Word installation
        queue.put({"ok": False, "error": str(exc)})


def _try_word_com_convert(source: Path, target: Path, timeout_seconds: int = 30) -> tuple[bool, str]:
    queue = mp.Queue()
    process = mp.Process(target=_word_com_convert_worker, args=(str(source.resolve()), str(target.resolve()), queue))
    process.start()
    process.join(timeout_seconds)
    if process.is_alive():
        process.terminate()
        process.join(5)
        return False, f"Word COM 转换超时（>{timeout_seconds}s），可能存在隐藏确认框或模板保护。"
    if not queue.empty():
        result = queue.get()
        return bool(result.get("ok")), str(result.get("error", ""))
    return target.exists(), "" if target.exists() else "Word COM 转换进程未返回结果。"


def convert_doc_to_docx(path: str | Path, output_dir: str | Path) -> Path:
    source = Path(path)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    errors: list[str] = []
    target = out_dir / f"{source.stem}.docx"
    converters = can_convert_doc()
    if converters["soffice"]:
        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(out_dir),
                    str(source),
                ],
                check=True,
            )
            converted = out_dir / f"{source.stem}.docx"
            if converted.exists():
                return converted
        except Exception as exc:
            errors.append(f"LibreOffice: {exc}")
    if converters["word_com"] and os.environ.get("THESIS_ASSISTANT_ENABLE_WORD_COM_CONVERT") == "1":
        ok, error = _try_word_com_convert(source, target)
        if ok and target.exists():
            return target
        if error:
            errors.append(f"Word COM: {error}")
    elif converters["word_com"]:
        errors.append("Word COM 已安装但默认不强制转换 .doc；如需尝试，设置 THESIS_ASSISTANT_ENABLE_WORD_COM_CONVERT=1。")
    detail = "；".join(errors) if errors else "当前未检测到可用的 soffice/winword 转换器。"
    raise RuntimeError(
        ".doc 文件需要先转换为 .docx；" + detail
    )


def ensure_docx(path: str | Path, output_dir: str | Path | None = None) -> Path:
    source = Path(path)
    if source.suffix.lower() == ".docx":
        return source
    if source.suffix.lower() == ".doc":
        return convert_doc_to_docx(source, output_dir or source.parent)
    raise ValueError(f"不支持的文件类型：{source.suffix}")


def load_document_info(path: str | Path) -> DocumentInfo:
    docx_path = ensure_docx(path)
    document = Document(str(docx_path))
    paragraphs = [
        ParagraphInfo(i, p.text.strip(), p.style.name if p.style else "")
        for i, p in enumerate(document.paragraphs)
        if p.text.strip()
    ]
    tables = [
        TableInfo(i, len(table.rows), len(table.columns))
        for i, table in enumerate(document.tables)
    ]
    picture_count = 0
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            picture_count += 1
    return DocumentInfo(
        path=docx_path,
        paragraphs=paragraphs,
        tables=tables,
        picture_count=picture_count,
        section_count=len(document.sections),
    )


def iter_paragraph_text(path: str | Path) -> Iterable[str]:
    document = Document(str(ensure_docx(path)))
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
