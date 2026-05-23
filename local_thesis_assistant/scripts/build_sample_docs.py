from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SAMPLES = ROOT / "data" / "samples"


def _set_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: int = 12, bold: bool = False):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def build_good_sample() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于本地论文助手的本科毕业设计说明书")
    _set_font(r, "黑体", "Times New Roman", 15, True)

    p = doc.add_paragraph()
    r = p.add_run("摘要")
    _set_font(r, "黑体", "Times New Roman", 14, True)

    p = doc.add_paragraph("本文设计并实现了一个本地论文助手，用于读取、检查和修正本科论文的 DOCX 文档格式。系统支持页边距、正文样式、标题层级、图表题注和引用顺序的基础校验，并可输出格式报告和规范化后的文档。")
    for run in p.runs:
        _set_font(run)

    p = doc.add_paragraph()
    r = p.add_run("1.1 需求分析")
    _set_font(r, "黑体", "Times New Roman", 14, True)

    p = doc.add_paragraph("系统面向学生、导师和论文辅助编辑场景，重点解决格式不一致、模板难复用和人工检查成本高的问题。")
    for run in p.runs:
        _set_font(run)

    p = doc.add_paragraph("图1.1 系统流程示意图")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        _set_font(run, size=10)

    p = doc.add_paragraph("表1.1 功能需求示例")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        _set_font(run, size=10)

    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 文档处理方法研究[J]. 计算机应用, 2024, 44(1): 1-8.")

    out = SAMPLES / "good_sample.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_bad_sample() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(4.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(4.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("第1章 绪论")
    _set_font(r, "Arial", "Arial", 10, False)

    p = doc.add_paragraph("这是一个故意错误的样例，用于测试格式检查器是否能发现字体、字号、页边距和题注问题。")
    for run in p.runs:
        _set_font(run, "Arial", "Arial", 9, False)

    p = doc.add_paragraph("图1.1 错误题注")
    for run in p.runs:
        _set_font(run, "Arial", "Arial", 9, False)

    p = doc.add_paragraph("图1.1 重复题注")
    for run in p.runs:
        _set_font(run, "Arial", "Arial", 9, False)

    p = doc.add_paragraph("[3][1] 引用顺序混乱")
    for run in p.runs:
        _set_font(run, "Arial", "Arial", 9, False)

    out = SAMPLES / "bad_sample.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main() -> None:
    good = build_good_sample()
    bad = build_bad_sample()
    print(good)
    print(bad)


if __name__ == "__main__":
    main()
