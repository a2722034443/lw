from __future__ import annotations

import argparse
import csv
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable


DEFAULT_CHAPTERS = [
    ("abstract", "摘要与关键词", 800, "研究目的、方法、系统实现、验证结果、结论。"),
    ("chapter_01", "第1章 绪论", 3000, "背景、意义、研究现状、问题分析、研究内容、技术路线。"),
    ("chapter_02", "第2章 相关技术与理论基础", 3000, "项目实际使用的技术、标准、方法和工具。"),
    ("chapter_03", "第3章 需求分析", 3000, "用户角色、功能需求、非功能需求、用例、可行性。"),
    ("chapter_04", "第4章 系统总体设计", 3200, "系统架构、模块划分、数据结构、接口、安全和部署。"),
    ("chapter_05", "第5章 系统详细设计与实现", 4300, "关键模块、核心流程、异常处理、页面或命令行实现。"),
    ("chapter_06", "第6章 系统测试与实验分析", 3400, "测试环境、测试用例、实验指标、结果和讨论。"),
    ("chapter_07", "第7章 工程管理与社会责任", 1500, "进度、成本、方案选择、隐私安全、法律伦理。"),
    ("chapter_08", "第8章 总结与展望", 1200, "完成工作、不足和后续改进。"),
]

LITERATURE_FIELDS = [
    "id",
    "title",
    "authors",
    "year",
    "source",
    "type",
    "doi",
    "url",
    "verified",
    "used_in_chapter",
    "key_point",
    "gbt7714",
]


@dataclass
class WorkspaceManifest:
    title: str
    thesis_type: str
    workspace: str
    created_by: str = "local_thesis_assistant.thesis_flow"
    policy: str = "正式正文、正式文献和实验结论必须联网核验；离线只能生成计划草稿。"


def _safe_name(value: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", value, flags=re.UNICODE).strip("_")
    return cleaned or "thesis_project"


def _write(path: Path, text: str, overwrite: bool = True) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists() and not overwrite:
        return path
    path.write_text(text, encoding="utf-8")
    return path


def _write_json(path: Path, data: object) -> Path:
    return _write(path, json.dumps(data, ensure_ascii=False, indent=2))


def _workspace(path: str | Path | None, title: str | None = None) -> Path:
    if path:
        return Path(path)
    return Path("local_thesis_assistant") / "outputs" / "thesis_flow" / _safe_name(title or "thesis_project")


def _ensure_dirs(root: Path) -> None:
    for name in [
        "literature",
        "chapters",
        "draft",
        "review",
        "assets",
        "logs",
    ]:
        (root / name).mkdir(parents=True, exist_ok=True)


def _load_manifest(root: Path) -> WorkspaceManifest:
    path = root / "workspace.json"
    if not path.exists():
        raise FileNotFoundError(f"未找到工作区清单：{path}，请先运行 init。")
    return WorkspaceManifest(**json.loads(path.read_text(encoding="utf-8")))


def _read_text_if_exists(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def cmd_init(args) -> int:
    root = _workspace(args.workspace, args.title)
    _ensure_dirs(root)
    manifest = WorkspaceManifest(args.title, args.type, str(root))
    _write_json(root / "workspace.json", asdict(manifest))
    _write_json(
        root / "topic_brief.json",
        {
            "title": args.title,
            "type": args.type,
            "system_name": "",
            "research_object": "",
            "users": [],
            "core_features": [],
            "data_sources": [],
            "experiment_needs": [],
            "risks": [],
            "innovation_points": [],
            "status": "draft",
        },
    )
    _write(
        root / "topic_brief.md",
        f"""# 选题评估：{args.title}

## 题目

{args.title}

## 研究对象

待联网调研后填写。

## 拟实现系统

待明确系统名称、用户角色和核心功能。

## 可行性

- 技术可行性：待评估。
- 数据可行性：待评估。
- 时间可行性：待评估。
- 论文可写性：待评估。

## 风险

- 文献未核验。
- 数据和实验尚未确认。
- 学校模板需优先确认。
""",
    )
    _write(root / "evidence_map.yaml", "evidence:\n")
    _write(root / "experiment_log.csv", "date,command,input,output,metric,value,notes\n")
    for key, title, _, goal in DEFAULT_CHAPTERS:
        _write(
            root / "chapters" / f"{key}.md",
            f"# {title}\n\n> 写作目标：{goal}\n\n> 状态：草稿，正式写作前必须补充证据链。\n\n",
            overwrite=False,
        )
    print(root)
    return 0


def cmd_research_standards(args) -> int:
    root = Path(args.workspace)
    manifest = _load_manifest(root)
    school = args.school or "未指定学校"
    text = f"""# 标准与模板调研报告

## 论文题目

{manifest.title}

## 学校

{school}

## 规则优先级

1. 学校或学院当年模板、格式要求、导师要求。
2. GB/T 7713.1-2025《信息与文献 编写规则 第1部分：学位论文》。
3. GB/T 7714-2015《信息与文献 参考文献著录规则》。
4. 软件项目类本科毕业设计常见结构。

## 必须联网核验的来源

- 国家标准全文公开系统：GB/T 7713.1-2025。
- 国家标准全文公开系统：GB/T 7714-2015。
- {school} 本科毕业设计说明书或论文格式要求。
- 软件工程项目型论文的需求分析、系统设计、系统测试章节写法。

## 当前状态

- 本报告是工作区标准记录骨架。
- 正式写作前必须补充检索日期、URL、截图或来源说明。
- 学校模板优先于通用标准。
"""
    if args.profile:
        text += f"\n## 已关联模板画像\n\n`{args.profile}`\n"
    _write(root / "standards_report.md", text)
    print(root / "standards_report.md")
    return 0


def cmd_research_literature(args) -> int:
    root = Path(args.workspace)
    _load_manifest(root)
    lit_dir = root / "literature"
    lit_dir.mkdir(parents=True, exist_ok=True)
    csv_path = lit_dir / "library.csv"
    if not csv_path.exists():
        with csv_path.open("w", encoding="utf-8-sig", newline="") as file:
            writer = csv.DictWriter(file, fieldnames=LITERATURE_FIELDS)
            writer.writeheader()
    queries = [
        f"{args.topic} 本科毕业设计 系统设计与实现",
        f"{args.topic} 文献综述",
        f"{args.topic} software system design implementation",
        f"{args.topic} document processing rule based formatting",
        "GB/T 7714-2015 参考文献 著录规则",
    ]
    _write(
        lit_dir / "research_queries.md",
        "# 文献检索计划\n\n"
        f"- 主题：{args.topic}\n"
        f"- 最少正式文献数量：{args.min}\n"
        "- 正式文献必须有 DOI、数据库页面、出版社页面或可信 URL。\n\n"
        + "\n".join(f"- `{item}`" for item in queries)
        + "\n",
    )
    _write(
        lit_dir / "verification.md",
        "# 文献核验记录\n\n| 编号 | 题名 | 来源 URL/DOI | 是否保留 | 使用章节 | 备注 |\n|---|---|---|---|---|---|\n",
        overwrite=False,
    )
    _write(
        lit_dir / "references_gbt7714.md",
        "# GB/T 7714-2015 参考文献草稿\n\n> 只有已联网核验的文献可以放入正式编号列表。\n",
        overwrite=False,
    )
    print(csv_path)
    return 0


def _outline_yaml(min_words: int, min_figures: int, min_tables: int) -> str:
    fig_targets = [0, 2, 4, 5, 6, 8, 5, 1, 0]
    table_targets = [0, 0, 1, 2, 1, 1, 2, 1, 0]
    total_fig = sum(fig_targets)
    total_tab = sum(table_targets)
    if min_figures > total_fig:
        fig_targets[5] += min_figures - total_fig
    if min_tables > total_tab:
        table_targets[6] += min_tables - total_tab
    lines = ["chapters:"]
    for index, (key, title, words, goal) in enumerate(DEFAULT_CHAPTERS):
        scaled_words = max(words, int(min_words * words / 24400))
        lines.extend(
            [
                f"  - id: {key}",
                f"    title: \"{title}\"",
                f"    target_words: {scaled_words}",
                f"    target_figures: {fig_targets[index]}",
                f"    target_tables: {table_targets[index]}",
                f"    writing_goal: \"{goal}\"",
                "    evidence_required: true" if key != "abstract" else "    evidence_required: true",
            ]
        )
    return "\n".join(lines) + "\n"


def cmd_build_outline(args) -> int:
    root = Path(args.workspace)
    manifest = _load_manifest(root)
    _write(root / "outline.yaml", _outline_yaml(args.min_words, args.min_figures, args.min_tables))
    _write(
        root / "writing_plan.md",
        f"""# 写作计划：{manifest.title}

## 硬性目标

- 正文净字数：不少于 {args.min_words}。
- 图：不少于 {args.min_figures}，必须有正文解释。
- 表：不少于 {args.min_tables}，默认三线表。
- 参考文献：正式条目必须联网核验。

## 写作顺序

1. 先补齐标准调研和文献库。
2. 完成项目设计和实验计划。
3. 建立证据链。
4. 分章写作。
5. 内容审查。
6. 导出 DOCX，并运行 audit/fix/link-references/verify。
""",
    )
    print(root / "outline.yaml")
    return 0


def cmd_design_project(args) -> int:
    root = Path(args.workspace)
    manifest = _load_manifest(root)
    _write(
        root / "project_spec.md",
        f"""# 项目设计说明：{manifest.title}

## 系统定位

待根据题目和选题评估填写系统目标、用户和使用场景。

## 用户角色

- 学生用户。
- 指导教师或审阅者。
- 系统维护者。

## 功能模块

| 模块 | 输入 | 处理 | 输出 | 对应论文章节 |
|---|---|---|---|---|
| 文档读取 | DOCX/DOC | 结构解析 | 文档信息 | 第5章 |
| 规则审查 | 文档结构 | 规则匹配 | 审查报告 | 第5章、第6章 |
| 保守修正 | 原 DOCX | 样式与段落修正 | 新 DOCX | 第5章 |
| 结果验证 | 修正文档 | 二次审查 | 验证报告 | 第6章 |

## 数据结构

待补充配置、报告、日志和实验记录结构。

## 图件需求

- 系统总体架构图。
- 文档处理流程图。
- 模块结构图。
- 审查与修正流程图。
- 测试流程图。
""",
    )
    _write(
        root / "project_architecture.yaml",
        "modules:\n  - name: 文档读取\n  - name: 规则审查\n  - name: 保守修正\n  - name: 验证报告\n",
    )
    _write(
        root / "figures_to_make.md",
        "# 图件清单\n\n- 系统总体架构图\n- 文档处理流程图\n- 用例图\n- 模块结构图\n- 测试流程图\n",
    )
    _write(
        root / "test_points.md",
        "# 测试点\n\n- 正常 DOCX 审查。\n- 错误格式 DOCX 修正。\n- 引用链接生成。\n- 表格三线表修正。\n- 缺失依赖诊断。\n",
    )
    print(root / "project_spec.md")
    return 0


def cmd_plan_experiment(args) -> int:
    root = Path(args.workspace)
    _load_manifest(root)
    _write(
        root / "experiment_plan.yaml",
        """experiments:
  - id: E01
    name: DOCX 格式错误检测
    input: data/samples/bad_sample.docx
    command: python -m local_thesis_assistant.thesis_assistant audit INPUT --profile PROFILE
    metrics: [finding_total, error_total, warning_total]
  - id: E02
    name: 保守修正前后对比
    input: data/samples/bad_sample.docx
    command: python -m local_thesis_assistant.thesis_assistant fix INPUT OUTPUT --profile PROFILE
    metrics: [before_total, after_total, after_error]
  - id: E03
    name: 参考文献内部跳转链接
    input: docx_with_references.docx
    command: python -m local_thesis_assistant.thesis_assistant link-references INPUT OUTPUT
    metrics: [linked_citations, skipped_citations]
""",
    )
    _write(
        root / "test_cases.md",
        "# 测试用例\n\n| 编号 | 测试目标 | 输入 | 预期结果 | 实际结果 |\n|---|---|---|---|---|\n| TC01 | 检查错误页边距 | 错误样例 DOCX | 报告 page_margin 错误 | 待运行 |\n| TC02 | 修正文档格式 | 错误样例 DOCX | 输出新 DOCX | 待运行 |\n| TC03 | 链接参考文献 | 含 [1] 引用 DOCX | 生成内部超链接 | 待运行 |\n",
    )
    _write(
        root / "result_tables.md",
        "# 实验结果表模板\n\n| 实验 | 修正前问题数 | 修正后问题数 | 说明 |\n|---|---:|---:|---|\n",
    )
    print(root / "experiment_plan.yaml")
    return 0


def cmd_write_section(args) -> int:
    root = Path(args.workspace)
    _load_manifest(root)
    chapter = str(args.chapter).zfill(2) if str(args.chapter).isdigit() else str(args.chapter)
    filename = "abstract.md" if chapter in {"0", "00", "abstract"} else f"chapter_{chapter}.md"
    target = root / "chapters" / filename
    if not target.exists():
        target = root / "chapters" / f"chapter_{chapter}.md"
    evidence_note = args.evidence or "evidence_map.yaml"
    content = _read_text_if_exists(target)
    if "## 证据链检查" not in content:
        content += f"\n## 证据链检查\n\n- 本章写作前必须核对 `{evidence_note}`。\n- 文献观点使用 `[n]` 引用。\n- 项目实现内容必须能追溯到 `project_spec.md` 或代码。\n- 实验结论必须能追溯到 `experiment_log.csv`。\n"
    _write(target, content)
    print(target)
    return 0


def _iter_chapter_files(root: Path) -> Iterable[Path]:
    for path in sorted((root / "chapters").glob("*.md")):
        yield path


def cmd_review_draft(args) -> int:
    root = Path(args.workspace)
    _load_manifest(root)
    findings: list[dict[str, str]] = []
    required = ["摘要", "绪论", "需求", "设计", "实现", "测试", "总结"]
    all_text = "\n".join(_read_text_if_exists(path) for path in _iter_chapter_files(root))
    for token in required:
        if token not in all_text:
            findings.append({"severity": "warning", "code": "missing_topic", "message": f"草稿中未明显出现：{token}"})
    for token in ["TODO", "待补", "这里写", "随便"]:
        if token in all_text:
            findings.append({"severity": "error", "code": "placeholder", "message": f"存在占位内容：{token}"})
    if "实验结果表明" in all_text and not (root / "experiment_log.csv").exists():
        findings.append({"severity": "error", "code": "experiment_without_log", "message": "写到实验结果但没有实验日志。"})
    citation_count = len(re.findall(r"\[[0-9]+\]", all_text))
    if citation_count == 0:
        findings.append({"severity": "warning", "code": "citation_absent", "message": "草稿中没有顺序编码引用。"})
    report = {
        "workspace": str(root),
        "findings": findings,
        "citation_count": citation_count,
        "chapter_files": [str(path) for path in _iter_chapter_files(root)],
    }
    _write_json(root / "review" / "content_review.json", report)
    lines = ["# 内容质量审查报告", "", f"- 引用数量：{citation_count}", ""]
    for item in findings:
        lines.append(f"- {item['severity'].upper()} `{item['code']}`：{item['message']}")
    if not findings:
        lines.append("- 未发现明显内容质量问题。")
    _write(root / "review" / "content_review.md", "\n".join(lines) + "\n")
    print(root / "review" / "content_review.md")
    return 1 if any(item["severity"] == "error" for item in findings) else 0


def _markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    from docx import Document

    document = Document()
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|"):
            document.add_paragraph(line)
        else:
            document.add_paragraph(line)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))


def cmd_export_docx(args) -> int:
    root = Path(args.workspace)
    _load_manifest(root)
    draft_dir = root / "draft"
    draft_dir.mkdir(parents=True, exist_ok=True)
    thesis_md = draft_dir / "thesis.md"
    combined = []
    for path in _iter_chapter_files(root):
        combined.append(path.read_text(encoding="utf-8").strip())
        combined.append("")
    _write(thesis_md, "\n\n".join(combined))
    raw_docx = Path(args.output) if args.output else draft_dir / "thesis.docx"
    _markdown_to_docx(thesis_md, raw_docx)
    final_docx = raw_docx
    if args.profile:
        from local_thesis_assistant.thesis_assistant.fixer import conservative_fix_docx

        fixed = raw_docx.with_name(raw_docx.stem + ".fixed.docx")
        conservative_fix_docx(raw_docx, fixed, args.profile)
        final_docx = fixed
    print(final_docx)
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="从题目到论文成稿的软件项目论文工作流工具")
    sub = parser.add_subparsers(dest="command", required=True)

    init = sub.add_parser("init", help="创建标准论文工作区")
    init.add_argument("--title", required=True)
    init.add_argument("--type", default="software-project")
    init.add_argument("--workspace")
    init.set_defaults(func=cmd_init)

    standards = sub.add_parser("research-standards", help="生成标准与模板调研记录")
    standards.add_argument("--workspace", required=True)
    standards.add_argument("--school", default="大连民族大学")
    standards.add_argument("--profile")
    standards.set_defaults(func=cmd_research_standards)

    literature = sub.add_parser("research-literature", help="生成文献检索与核验工作表")
    literature.add_argument("--workspace", required=True)
    literature.add_argument("--topic", required=True)
    literature.add_argument("--min", type=int, default=20)
    literature.set_defaults(func=cmd_research_literature)

    outline = sub.add_parser("build-outline", help="生成论文大纲与写作计划")
    outline.add_argument("--workspace", required=True)
    outline.add_argument("--profile")
    outline.add_argument("--min-words", type=int, default=20000)
    outline.add_argument("--min-figures", type=int, default=40)
    outline.add_argument("--min-tables", type=int, default=6)
    outline.set_defaults(func=cmd_build_outline)

    design = sub.add_parser("design-project", help="生成项目设计文档")
    design.add_argument("--workspace", required=True)
    design.set_defaults(func=cmd_design_project)

    experiment = sub.add_parser("plan-experiment", help="生成实验与测试计划")
    experiment.add_argument("--workspace", required=True)
    experiment.set_defaults(func=cmd_plan_experiment)

    writer = sub.add_parser("write-section", help="为章节草稿补写证据链提示")
    writer.add_argument("--workspace", required=True)
    writer.add_argument("--chapter", required=True)
    writer.add_argument("--evidence", default="evidence_map.yaml")
    writer.set_defaults(func=cmd_write_section)

    review = sub.add_parser("review-draft", help="审查章节草稿内容质量")
    review.add_argument("--workspace", required=True)
    review.set_defaults(func=cmd_review_draft)

    export = sub.add_parser("export-docx", help="汇总章节草稿并导出 DOCX")
    export.add_argument("--workspace", required=True)
    export.add_argument("--output")
    export.add_argument("--profile")
    export.set_defaults(func=cmd_export_docx)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
