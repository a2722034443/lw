import argparse
import shutil
import subprocess
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")


def extract_with_pandoc(docx_path: Path, track_changes: str) -> str | None:
    if not shutil.which("pandoc"):
        return None
    cmd = ["pandoc", f"--track-changes={track_changes}", str(docx_path), "-t", "markdown"]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "pandoc 提取失败")
    return result.stdout


def extract_with_python_docx(docx_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法降级提取 DOCX 文本") from exc

    doc = Document(str(docx_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"\n[表格 {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="提取 DOCX 文本，优先使用 pandoc，缺失时使用 python-docx。")
    parser.add_argument("docx", type=Path)
    parser.add_argument("-o", "--output", type=Path)
    parser.add_argument("--track-changes", choices=["all", "accept", "reject"], default="all")
    args = parser.parse_args()

    if not args.docx.exists():
        parser.error(f"文件不存在：{args.docx}")
    text = extract_with_pandoc(args.docx, args.track_changes)
    if text is None:
        text = extract_with_python_docx(args.docx)
    if args.output:
        args.output.write_text(text, encoding="utf-8")
    else:
        sys.stdout.write(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
