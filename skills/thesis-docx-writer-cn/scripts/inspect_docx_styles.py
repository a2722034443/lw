import argparse
import json
import sys
import zipfile
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")


def inspect_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法检查样式") from exc

    doc = Document(str(path))
    styles = []
    for style in doc.styles:
        styles.append({"name": style.name, "type": str(style.type)})

    paragraphs = []
    for index, para in enumerate(doc.paragraphs[:80], start=1):
        if para.text.strip():
            paragraphs.append({"index": index, "style": para.style.name if para.style else "", "text": para.text[:120]})

    sections = []
    for index, section in enumerate(doc.sections, start=1):
        sections.append(
            {
                "index": index,
                "page_width": section.page_width.pt,
                "page_height": section.page_height.pt,
                "top_margin": section.top_margin.pt,
                "bottom_margin": section.bottom_margin.pt,
                "left_margin": section.left_margin.pt,
                "right_margin": section.right_margin.pt,
            }
        )

    rels_count = 0
    image_count = 0
    with zipfile.ZipFile(path, "r") as zf:
        names = zf.namelist()
        rels_count = len([n for n in names if n.endswith(".rels")])
        image_count = len([n for n in names if n.startswith("word/media/")])

    return {
        "file": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "relationship_file_count": rels_count,
        "image_count": image_count,
        "sections": sections,
        "paragraph_samples": paragraphs,
        "styles": styles,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="检查 DOCX 模板样式、段落、表格、图片和页面设置。")
    parser.add_argument("docx", type=Path)
    parser.add_argument("-o", "--output", type=Path)
    args = parser.parse_args()

    if not args.docx.exists():
        parser.error(f"文件不存在：{args.docx}")
    report = inspect_docx(args.docx)
    text = json.dumps(report, ensure_ascii=False, indent=2)
    if args.output:
        args.output.write_text(text, encoding="utf-8")
    else:
        print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
