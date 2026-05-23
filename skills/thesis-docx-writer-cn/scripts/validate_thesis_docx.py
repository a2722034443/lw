import argparse
import json
import re
import sys
import zipfile
from collections import defaultdict
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
FIG_RE = re.compile(r"^\s*图\s*([0-9]+)[\.-]([0-9]+)")
TABLE_RE = re.compile(r"^\s*表\s*([0-9]+)[\.-]([0-9]+)")
CITE_RE = re.compile(r"\[([0-9]+)\]")
REF_LINE_RE = re.compile(r"^\s*\[([0-9]+)\]\s*")


def non_space_chars(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def read_document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")


def media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        return len([name for name in zf.namelist() if name.startswith("word/media/")])


def field_codes(xml: str) -> list[str]:
    root = ET.fromstring(xml)
    codes = []
    for node in root.findall(".//w:instrText", NS):
        if node.text:
            codes.append(" ".join(node.text.split()))
    return codes


def caption_numbers(paragraphs: list[str], pattern: re.Pattern[str]) -> list[tuple[int, int]]:
    found = []
    for text in paragraphs:
        m = pattern.search(text)
        if m:
            found.append((int(m.group(1)), int(m.group(2))))
    return found


def duplicate_or_gap(numbers: list[tuple[int, int]]) -> dict[str, list[str]]:
    by_chapter: dict[int, list[int]] = defaultdict(list)
    for chapter, idx in numbers:
        by_chapter[chapter].append(idx)
    duplicates = []
    gaps = []
    for chapter, items in sorted(by_chapter.items()):
        seen = set()
        for item in items:
            if item in seen:
                duplicates.append(f"{chapter}.{item}")
            seen.add(item)
        ordered = sorted(seen)
        if ordered:
            expected = list(range(ordered[0], ordered[-1] + 1))
            missing = [x for x in expected if x not in seen]
            gaps.extend([f"{chapter}.{x}" for x in missing])
    return {"duplicates": duplicates, "gaps": gaps}


def first_seen_citations(body: str) -> list[int]:
    seen = []
    for raw in CITE_RE.findall(body):
        n = int(raw)
        if n not in seen:
            seen.append(n)
    return seen


def reference_numbers(full_text: str) -> list[int]:
    start = full_text.rfind("参考文献")
    if start < 0:
        return []
    return [int(x) for x in re.findall(r"(?m)^\s*\[([0-9]+)\]\s*", full_text[start:])]


def paragraph_format_summary(doc: Document, sample_limit: int = 80) -> dict[str, int]:
    summary = {
        "normal_like_sample": 0,
        "first_line_indent_present": 0,
        "line_spacing_present": 0,
        "font_size_present": 0,
        "east_asia_font_present": 0,
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            continue
        if FIG_RE.search(text) or TABLE_RE.search(text) or REF_LINE_RE.search(text):
            continue
        summary["normal_like_sample"] += 1
        if p.paragraph_format.first_line_indent is not None:
            summary["first_line_indent_present"] += 1
        if p.paragraph_format.line_spacing is not None:
            summary["line_spacing_present"] += 1
        for run in p.runs:
            if run.font.size is not None:
                summary["font_size_present"] += 1
                break
        for run in p.runs:
            rpr = run._element.rPr
            if rpr is not None and rpr.rFonts is not None and rpr.rFonts.get(qn("w:eastAsia")):
                summary["east_asia_font_present"] += 1
                break
        if summary["normal_like_sample"] >= sample_limit:
            break
    return summary


def qn(tag: str) -> str:
    prefix, name = tag.split(":")
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return f"{{{namespaces[prefix]}}}{name}"


def main() -> int:
    parser = argparse.ArgumentParser(description="检查中文项目论文 DOCX 的字数、图表、引用、题注和基础版式。")
    parser.add_argument("docx", type=Path, help="待检查的 docx 文件")
    parser.add_argument("--min-chars", type=int, default=20000, help="正文最低净字数")
    parser.add_argument("--min-figures", type=int, default=40, help="最低图片数量")
    parser.add_argument("--min-tables", type=int, default=6, help="最低表格数量")
    parser.add_argument("--require-field-refs", action="store_true", help="要求存在 Word REF 交叉引用域")
    parser.add_argument("--json", action="store_true", help="输出 JSON")
    args = parser.parse_args()

    doc = Document(args.docx)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    body_end = full_text.rfind("参考文献")
    body = full_text if body_end < 0 else full_text[:body_end]
    body_chars = non_space_chars(body)
    embedded_media = media_count(args.docx)
    figure_numbers = caption_numbers(paragraphs, FIG_RE)
    table_numbers = caption_numbers(paragraphs, TABLE_RE)
    references = reference_numbers(full_text)
    citations = first_seen_citations(body)
    xml = read_document_xml(args.docx)
    codes = field_codes(xml)
    ref_fields = [c for c in codes if re.search(r"\bREF\b|\bPAGEREF\b", c)]
    seq_fields = [c for c in codes if re.search(r"\bSEQ\b", c)]
    toc_fields = [c for c in codes if re.search(r"\bTOC\b", c)]

    issues = []
    warnings = []
    if body_chars < args.min_chars:
        issues.append(f"正文净字数不足：{body_chars} < {args.min_chars}")
    if len(figure_numbers) < args.min_figures:
        issues.append(f"图题数量不足：{len(figure_numbers)} < {args.min_figures}")
    if embedded_media < args.min_figures:
        warnings.append(f"嵌入媒体数量少于目标图片数：{embedded_media} < {args.min_figures}")
    if len(doc.tables) < args.min_tables and len(table_numbers) < args.min_tables:
        issues.append(f"表格数量不足：docx表格{len(doc.tables)}个，表题{len(table_numbers)}个，目标{args.min_tables}个")
    if references:
        expected = list(range(references[0], references[0] + len(references)))
        if references != expected:
            issues.append("参考文献编号不连续")
    if citations:
        if citations != sorted(citations):
            issues.append("正文首次引用顺序不是递增顺序")
        if references:
            missing = [n for n in citations if n not in references]
            unused = [n for n in references if n not in citations]
            if missing:
                issues.append("正文引用在参考文献表中不存在：" + ",".join(map(str, missing)))
            if unused:
                issues.append("参考文献未在正文中引用：" + ",".join(map(str, unused)))
    if not seq_fields:
        warnings.append("未检测到 SEQ 题注编号域；可能使用的是静态图表编号")
    if not ref_fields:
        message = "未检测到 REF/PAGEREF 交叉引用域；可能使用的是静态交叉引用"
        if args.require_field_refs:
            issues.append(message)
        else:
            warnings.append(message)
    if not toc_fields:
        warnings.append("未检测到 TOC 目录域；目录可能是静态文本")

    figure_sequence = duplicate_or_gap(figure_numbers)
    table_sequence = duplicate_or_gap(table_numbers)
    if figure_sequence["duplicates"]:
        issues.append("图号重复：" + ",".join(figure_sequence["duplicates"]))
    if figure_sequence["gaps"]:
        warnings.append("图号存在跳号：" + ",".join(figure_sequence["gaps"]))
    if table_sequence["duplicates"]:
        issues.append("表号重复：" + ",".join(table_sequence["duplicates"]))
    if table_sequence["gaps"]:
        warnings.append("表号存在跳号：" + ",".join(table_sequence["gaps"]))

    result = {
        "file": str(args.docx),
        "body_chars": body_chars,
        "figure_captions": len(figure_numbers),
        "embedded_media": embedded_media,
        "docx_tables": len(doc.tables),
        "table_captions": len(table_numbers),
        "reference_count": len(references),
        "first_seen_citations": citations,
        "field_codes": {
            "SEQ": len(seq_fields),
            "REF_OR_PAGEREF": len(ref_fields),
            "TOC": len(toc_fields),
        },
        "paragraph_format_sample": paragraph_format_summary(doc),
        "issues": issues,
        "warnings": warnings,
    }
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"文件：{args.docx}")
        print(f"正文净字数：{body_chars}")
        print(f"图题数量：{len(figure_numbers)}，嵌入媒体：{embedded_media}")
        print(f"表题数量：{len(table_numbers)}，docx表格：{len(doc.tables)}")
        print(f"参考文献：{len(references)}，正文首次引用：{citations}")
        print(f"Word域：SEQ={len(seq_fields)}，REF/PAGEREF={len(ref_fields)}，TOC={len(toc_fields)}")
        for item in warnings:
            print("提醒：" + item)
        for item in issues:
            print("错误：" + item)
        if not issues:
            print("通过硬性检查")
    return 1 if issues else 0


if __name__ == "__main__":
    raise SystemExit(main())
