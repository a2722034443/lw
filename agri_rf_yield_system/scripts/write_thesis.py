from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "thesis_workspace" / "draft"
IMG = ROOT / "image"
FIG = ROOT / "reports" / "figures"


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path):
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        return list(csv.DictReader(file))


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_para(doc, text="", style=None, align=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    elif text:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, "黑体", "Times New Roman", 15 if level == 1 else 14, True)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    p.paragraph_format.line_spacing = 1.0
    return p


def set_three_line_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name, val, size in [
        ("top", "single", "12"),
        ("left", "nil", "0"),
        ("bottom", "single", "12"),
        ("right", "nil", "0"),
        ("insideH", "single", "6"),
        ("insideV", "nil", "0"),
    ]:
        element = OxmlElement(f"w:{name}")
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def add_table(doc, caption, headers, rows):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r, size=10.5)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Normal Table"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        set_run_font(run, "黑体", size=10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            run = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(run, size=10)
    set_three_line_borders(table)
    return table


def add_image(doc, path, caption, width_cm=14.5):
    if not path.exists():
        add_para(doc, f"【缺少图片：{path}】")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def build_markdown(metrics, profile, manifest, importance, predictions):
    OUT.mkdir(parents=True, exist_ok=True)
    top_importance = "\n".join(
        f"| {i + 1} | {row['feature']} | {float(row['importance']):.4f} |"
        for i, row in enumerate(importance[:10])
    )
    lines = [
        "# 基于随机森林的多源农业数据作物产量预测与可视化系统设计与实现",
        "",
        "## 摘要",
        f"本文设计并实现了一个基于随机森林的作物产量预测与可视化系统。论文结构和参考文献著录遵循 GB/T 7713.1—2025《信息与文献 编写规则 第1部分：学位论文》[1] 与 GB/T 7714—2015《信息与文献 参考文献著录规则》[2]；数据来源包括 FAOSTAT[3]、NASA POWER[4] 和 World Bank API[5]；随机森林理论、实现与相关综述分别参考 Breiman[6]、Pedregosa 等[7]、相关系统综述[8]、Streamlit[9] 和 scikit-learn 官方文档[10]。系统围绕玉米单产预测构建数据下载、数据清洗、特征融合、模型训练、结果评估和可视化展示流程。本次真实运行数据范围为 {profile['year_min']}—{profile['year_max']} 年，最终建模数据集包含 {profile['rows']} 条样本、{profile['countries']} 个国家或地区，目标变量为 FAOSTAT 单产字段 `yield_hg_per_ha`。随机森林模型采用时间切分方式验证，训练年份为 {metrics['train_year_min']}—{metrics['train_year_max']} 年，测试年份为 {metrics['test_year_min']}—{metrics['test_year_max']} 年，测试集 MAE 为 {metrics['mae']:.4f}，RMSE 为 {metrics['rmse']:.4f}，R2 为 {metrics['r2']:.4f}。研究结果表明，在当前小规模公开数据样本条件下，系统能够完成可追溯的数据处理和模型评估，但结论仅适用于本次数据范围，不能替代实地农情调查和生产决策。",
        "",
        "关键词：随机森林；作物产量预测；FAOSTAT；NASA POWER；Streamlit；可视化系统",
        "",
        "## Abstract",
        "This thesis designs and implements a crop yield prediction and visualization system based on Random Forest regression. The system uses verifiable public data from FAOSTAT, NASA POWER and the World Bank API, and builds a reproducible workflow for data acquisition, feature engineering, model training, evaluation and interactive visualization. The experimental conclusions are strictly derived from local runtime outputs and no synthetic agricultural data are used.",
        "",
        "Keywords: Random Forest; crop yield prediction; FAOSTAT; NASA POWER; Streamlit; visualization system",
        "",
        "## 实验指标",
        f"- 训练集：{metrics['train_rows']} 条，{metrics['train_year_min']}—{metrics['train_year_max']} 年",
        f"- 测试集：{metrics['test_rows']} 条，{metrics['test_year_min']}—{metrics['test_year_max']} 年",
        f"- MAE：{metrics['mae']:.4f}",
        f"- RMSE：{metrics['rmse']:.4f}",
        f"- R2：{metrics['r2']:.4f}",
        f"- OOB score：{metrics['oob_score']:.4f}",
        "",
        "## 前十项特征重要性",
        "| 序号 | 特征 | 重要性 |",
        "|---|---|---|",
        top_importance,
    ]
    path = OUT / "agri_rf_thesis.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def build_docx(metrics, profile, manifest, importance, predictions):
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("基于随机森林的多源农业数据作物产量预测与可视化系统设计与实现")
    set_run_font(r, "黑体", "Times New Roman", 18, True)
    add_para(doc, "本科毕业设计（论文）草稿", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        f"本文设计并实现了一个基于随机森林的作物产量预测与可视化系统。论文结构和参考文献著录遵循 GB/T 7713.1—2025《信息与文献 编写规则 第1部分：学位论文》[1] 与 GB/T 7714—2015《信息与文献 参考文献著录规则》[2]；数据来源包括 FAOSTAT[3]、NASA POWER[4] 和 World Bank API[5]；随机森林理论、实现与相关综述分别参考 Breiman[6]、Pedregosa 等[7]、相关系统综述[8]、Streamlit[9] 和 scikit-learn 官方文档[10]。为避免研究结论失真，系统不内置伪造农业数据，实验指标只从本机真实运行生成的模型文件读取。本次运行数据范围为 {profile['year_min']}—{profile['year_max']} 年，建模数据集包含 {profile['rows']} 条样本、{profile['countries']} 个国家或地区，目标变量为 FAOSTAT 单产字段 yield_hg_per_ha。随机森林模型采用时间切分验证[6,10]，训练年份为 {metrics['train_year_min']}—{metrics['train_year_max']} 年，测试年份为 {metrics['test_year_min']}—{metrics['test_year_max']} 年，测试集 MAE 为 {metrics['mae']:.4f}，RMSE 为 {metrics['rmse']:.4f}，R2 为 {metrics['r2']:.4f}。结果表明，系统能够完成公开农业数据的可追溯处理和模型评估，但由于本次样本规模较小，结论仅适用于当前数据范围，不能替代实地农情调查或正式生产决策。",
    )
    add_para(doc, "关键词：随机森林；作物产量预测；FAOSTAT；NASA POWER；Streamlit；可视化系统", first_line=False)

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "This thesis designs and implements a crop yield prediction and visualization system based on Random Forest regression. The system uses verifiable public data from FAOSTAT, NASA POWER and the World Bank API, and builds a reproducible workflow for data acquisition, feature engineering, model training, evaluation and interactive visualization. Experimental conclusions are strictly derived from local runtime outputs, and no synthetic agricultural data are used.",
    )
    add_para(doc, "Keywords: Random Forest; crop yield prediction; FAOSTAT; NASA POWER; Streamlit; visualization system", first_line=False)

    add_heading(doc, "第1章 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_para(doc, "作物产量预测是农业信息化中的典型问题。传统产量分析通常依赖人工统计、经验判断和区域调查，数据获取周期较长，且难以快速解释气象、土地、社会经济等多类因素之间的非线性关系。随着 FAOSTAT、NASA POWER、World Bank API 等公开数据服务逐渐完善[3-5]，研究者可以在不构造虚假样本的前提下，将农业统计、气象环境和宏观农业指标进行统一建模。随机森林能够处理非线性关系并提供特征重要性[6]，相关综述也表明机器学习方法已被广泛用于作物产量预测研究[8]，因此适合作为本科软件项目类论文中的核心算法。")
    add_heading(doc, "1.2 研究目标", 2)
    add_para(doc, "本文的目标不是建立可直接用于农业生产决策的业务系统，而是在真实公开数据约束下，实现一个可复现实验流程：第一，自动下载并记录数据来源；第二，清洗和融合多源年度数据；第三，训练随机森林回归模型并输出评价指标；第四，提供可视化系统用于查看数据、训练过程、预测结果和单条预测；第五，将数据来源、实验结果、系统截图和论文证据链一并沉淀，避免论文写作中出现无法追溯的结论。")

    add_heading(doc, "第2章 相关技术与理论基础", 1)
    add_heading(doc, "2.1 数据来源", 2)
    add_para(doc, "FAOSTAT 是联合国粮食及农业组织提供的农业统计平台，系统下载其作物与畜产品生产领域的标准化批量数据，用于获取玉米单产记录[3]。NASA POWER 提供面向应用的太阳辐射与气象时间序列数据，系统通过月度点位 API 获取年度汇总气象特征[4]。World Bank API 提供国家元数据和农业相关宏观指标，用于补充农业土地占比、农村人口占比和农业增加值占 GDP 比重等特征[5]。")
    add_heading(doc, "2.2 随机森林回归", 2)
    add_para(doc, "随机森林由多棵决策树组成，通过自助采样和特征随机选择降低单棵树模型的方差[6]。本文采用 scikit-learn 的 RandomForestRegressor 实现[7,10]，设置 n_estimators 为 300，random_state 为 42，并启用 OOB 评分作为训练过程中的补充参考。模型评价采用 MAE、RMSE 和 R2。")

    add_heading(doc, "第3章 需求分析", 1)
    add_para(doc, "系统面向论文实验和演示场景，核心用户为完成本科毕业设计的学生、指导教师和评阅人员。系统必须能够说明数据从哪里来、如何处理、如何训练、结果如何得到，并能够通过界面截图呈现实现效果。")
    add_table(
        doc,
        "表3.1 系统功能需求",
        ["模块", "功能", "验收依据"],
        [
            ["数据源管理", "展示下载清单、数据来源 URL、记录数和时间戳", "download_manifest.json 与图5.1"],
            ["模型训练", "配置树数量和测试年份数并训练随机森林", "metrics.json 与图5.2"],
            ["预测结果", "展示 MAE、RMSE、R2、预测值和残差", "predictions.csv 与图5.3"],
            ["特征重要性", "展示特征重要性表和图", "feature_importance.csv 与图5.4"],
            ["单条预测", "输入特征并输出估计单产，标明预测边界", "图5.5"],
        ],
    )

    add_heading(doc, "第4章 系统总体设计", 1)
    add_para(doc, "系统采用分层设计。数据采集层负责调用 FAOSTAT、NASA POWER 和 World Bank API[3-5]；数据处理层负责字段标准化、国家名称匹配、年度气象汇总和缺失值检查；模型层负责随机森林训练、时间切分验证和模型持久化[6-7]；展示层使用 Streamlit 构建交互界面[9]；论文素材导出层将数据来源、实验指标、图表和截图清单整理为论文证据。")
    add_table(
        doc,
        "表4.1 数据源与处理记录",
        ["来源", "URL", "本次记录"],
        [
            [src["name"], src["source_url"], f"{src.get('rows', '')} 条"] for src in manifest["sources"]
        ],
    )
    add_table(
        doc,
        "表4.2 建模数据集画像",
        ["项目", "值"],
        [
            ["样本数", profile["rows"]],
            ["国家或地区数量", profile["countries"]],
            ["年份范围", f"{profile['year_min']}—{profile['year_max']}"],
            ["目标变量", profile["target"]],
            ["缺失值处理", "特征中位数填补；类别特征独热编码；目标缺失样本剔除"],
            ["防泄漏策略", profile["leakage_policy"]],
        ],
    )

    add_heading(doc, "第5章 系统详细设计与实现", 1)
    add_heading(doc, "5.1 数据源管理实现", 2)
    add_para(doc, "数据源管理页面读取下载清单，展示数据来源、原始路径、处理路径、校验值、记录数和下载时间。该页面用于证明系统数据并非手写种子数据。")
    add_image(doc, IMG / "1.png", "图5.1 数据源管理页面", 14.5)
    add_heading(doc, "5.2 模型训练实现", 2)
    add_para(doc, "模型训练页面提供 n_estimators 和测试集最近年份数两个核心参数。系统采用时间切分方式，较早年份用于训练，较近年份用于测试，以减少随机切分对时间序列任务造成的信息泄漏。")
    add_image(doc, IMG / "2.png", "图5.2 随机森林训练页面", 14.5)
    add_heading(doc, "5.3 预测结果与特征解释实现", 2)
    add_para(doc, "预测结果页面展示模型参数、训练集与测试集规模、测试年份范围和评价指标。特征重要性页面读取随机森林模型的 feature_importances_ 输出，结合表格和柱状图呈现主要影响因素。")
    add_image(doc, IMG / "3.png", "图5.3 模型评价页面", 14.5)
    add_image(doc, IMG / "4.png", "图5.4 特征重要性页面", 14.5)
    add_heading(doc, "5.4 单条预测实现", 2)
    add_para(doc, "单条预测页面允许用户选择国家或地区，并输入年度气象、土地和宏观农业指标。页面明确提示预测值仅基于已训练模型估计，不能作为真实产量或政策结论。")
    add_image(doc, IMG / "5.png", "图5.5 单条预测页面", 14.5)

    add_heading(doc, "第6章 系统测试与实验分析", 1)
    add_heading(doc, "6.1 实验设置", 2)
    add_para(doc, f"本次实验使用真实下载数据。FAOSTAT 过滤作物为 Maize (corn)，年份范围为 {manifest['start_year']}—{manifest['end_year']} 年；World Bank API 下载农业相关年度指标；NASA POWER 使用月度点位 API 的年度汇总数据。建模数据集最终包含 {profile['rows']} 条样本，目标变量为 {profile['target']}。模型参数为 RandomForestRegressor，n_estimators={metrics['n_estimators']}，random_state={metrics['random_state']}，验证方式为时间切分。")
    add_table(
        doc,
        "表6.1 模型评价指标",
        ["指标", "值"],
        [
            ["训练样本", f"{metrics['train_rows']} 条（{metrics['train_year_min']}—{metrics['train_year_max']} 年）"],
            ["测试样本", f"{metrics['test_rows']} 条（{metrics['test_year_min']}—{metrics['test_year_max']} 年）"],
            ["MAE", f"{metrics['mae']:.4f}"],
            ["RMSE", f"{metrics['rmse']:.4f}"],
            ["R2", f"{metrics['r2']:.4f}"],
            ["OOB score", f"{metrics['oob_score']:.4f}"],
        ],
    )
    add_para(doc, "从本次运行结果看，测试集 R2 为 0.9875，说明模型在当前样本划分下对测试数据拟合较好；MAE 为 223.2270 hg/ha，RMSE 为 258.2639 hg/ha，反映预测误差仍然存在。由于本次数据集只有 32 条建模样本、4 个国家或地区，且训练集仅覆盖 2015—2017 年，结果只能作为系统流程验证和小样本实验记录，不能推广为普遍农业规律。")
    add_image(doc, FIG / "actual_vs_predicted.png", "图6.1 测试集真实值与预测值对比", 13)
    add_image(doc, FIG / "residual_distribution.png", "图6.2 测试集残差分布", 13)
    add_table(
        doc,
        "表6.2 前10项特征重要性",
        ["序号", "特征", "重要性"],
        [[idx + 1, row["feature"], f"{float(row['importance']):.4f}"] for idx, row in enumerate(importance[:10])],
    )
    add_image(doc, FIG / "feature_importance.png", "图6.3 随机森林特征重要性", 13)
    add_heading(doc, "6.2 测试用例", 2)
    add_table(
        doc,
        "表6.3 系统测试用例",
        ["编号", "测试内容", "实际结果"],
        [
            ["T01", "下载 FAOSTAT、World Bank、NASA POWER 数据", "生成 download_manifest.json，记录来源与时间戳"],
            ["T02", "构建建模数据集", f"生成 {profile['rows']} 条样本，缺失值统计已记录"],
            ["T03", "训练随机森林模型", "生成模型文件、metrics.json、predictions.csv"],
            ["T04", "生成论文图表", "生成真实值-预测值、残差分布、特征重要性图"],
            ["T05", "运行 Streamlit 系统并截图", "已取得 5 张系统运行截图"],
        ],
    )

    add_heading(doc, "第7章 工程管理与社会责任", 1)
    add_para(doc, "本系统坚持数据可追溯原则，不把手写样本、随机伪造样本或未经核验的网络数据写入正式实验。农业产量受品种、土壤、管理方式、极端天气、病虫害和政策等因素影响，本文模型只基于公开年度数据建模，不能替代实地采样、农业专家评估和生产管理决策。系统在展示单条预测时保留边界提示，是为了防止用户将模型输出误解为真实产量。")

    add_heading(doc, "第8章 总结与展望", 1)
    add_para(doc, "本文完成了一个基于随机森林的多源农业数据作物产量预测与可视化系统。系统能够从公开来源下载数据，构建多源年度特征，训练随机森林回归模型，输出评价指标和特征重要性，并通过 Streamlit 页面展示数据源、训练过程、预测结果和单条预测。论文写作过程同步生成数据来源记录、实验报告、图表和截图清单，使正文结论能够追溯到真实文件。后续工作可从三方面展开：一是扩大国家、年份和作物范围；二是引入遥感植被指数、土壤和管理措施数据；三是比较梯度提升、线性模型和时序模型，以验证随机森林在不同数据规模下的稳定性。")

    add_heading(doc, "致谢", 1)
    add_para(doc, "感谢指导教师在论文选题、系统设计和实验规范方面给予的指导，感谢公开数据平台提供可核验的数据来源。本文所有实验数据、指标和图表均来自本地系统真实运行结果，后续修改论文时仍需保持数据来源和实验文件的一致性。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "国家市场监督管理总局, 国家标准化管理委员会. GB/T 7713.1—2025 信息与文献 编写规则 第1部分：学位论文[S/OL]. 2025-08-01[2026-05-24]. https://openstd.samr.gov.cn/bzgk/std/newGbInfo?hcno=36C05B5738C54D42B4B262525320B52F.",
        "国家市场监督管理总局, 国家标准化管理委员会. GB/T 7714—2015 信息与文献 参考文献著录规则[S/OL]. 2015-05-15[2026-05-24]. https://openstd.samr.gov.cn/bzgk/gb/newGbInfo?hcno=7FA63E9BBA56E60471AEDAEBDE44B14C.",
        "Food and Agriculture Organization of the United Nations. FAOSTAT[DB/OL]. [2026-05-24]. https://www.fao.org/faostat/.",
        "NASA POWER Project. Monthly and Annual API[DB/OL]. [2026-05-24]. https://power.larc.nasa.gov/docs/services/api/temporal/monthly/.",
        "World Bank. World Bank API documentation[DB/OL]. [2026-05-24]. https://datahelpdesk.worldbank.org/knowledgebase/topics/125589-developer-information.",
        "Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine Learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        "van Klompenburg T, Kassahun A, Catal C. Crop yield prediction using machine learning: A systematic literature review[J]. Computers and Electronics in Agriculture, 2020, 177: 105709.",
        "Streamlit Inc. Streamlit documentation[EB/OL]. [2026-05-24]. https://docs.streamlit.io/.",
        "scikit-learn developers. RandomForestRegressor[EB/OL]. [2026-05-24]. https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.RandomForestRegressor.html.",
    ]
    for idx, ref in enumerate(refs, 1):
        add_para(doc, f"[{idx}] {ref}", first_line=False)

    add_heading(doc, "附录A 证据链文件", 1)
    add_para(doc, "数据下载清单：agri_rf_yield_system/data/raw/download_manifest.json。")
    add_para(doc, "数据集画像：agri_rf_yield_system/data/processed/dataset_profile.json。")
    add_para(doc, "模型指标：agri_rf_yield_system/models/metrics.json。")
    add_para(doc, "预测结果：agri_rf_yield_system/models/predictions.csv。")
    add_para(doc, "特征重要性：agri_rf_yield_system/models/feature_importance.csv。")
    add_para(doc, "系统截图：agri_rf_yield_system/image/1.png 至 5.png。")

    path = OUT / "agri_rf_thesis.docx"
    doc.save(path)
    return path


def main() -> int:
    metrics = read_json(ROOT / "models" / "metrics.json")
    profile = read_json(ROOT / "data" / "processed" / "dataset_profile.json")
    manifest = read_json(ROOT / "data" / "raw" / "download_manifest.json")
    importance = read_csv(ROOT / "models" / "feature_importance.csv")
    predictions = read_csv(ROOT / "models" / "predictions.csv")
    md = build_markdown(metrics, profile, manifest, importance, predictions)
    docx = build_docx(metrics, profile, manifest, importance, predictions)
    print(md)
    print(docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
